"""
================================================================================
SYMBOLIC ENTROPY (SE) ANALYSIS TOOL - Streamlined Production Version
================================================================================

A computational framework for quantifying meaning density in literary texts
through the combination of Shannon entropy (lexical diversity) and KL divergence
(archetypal motif concentration).

THEORETICAL FOUNDATION:
    SE = (H , Σ)
    
    Where:
        H = Shannon entropy (bits/token) - measures lexical unpredictability
        Σ = KL divergence (bits/token) - measures motif clustering beyond baseline
        
    This framework extends Shannon's 1948 information theory by adding semantic
    content measurement to purely statistical analysis.

FALSIFIABILITY:
    SE is scientifically falsifiable. Empirical testing shows that when text 
    is randomized (destroying narrative structure while preserving vocabulary), 
    Σ collapses to near-zero, proving SE measures structure not word frequency.
    
    Note: For large texts (>50k tokens), global baseline converges to shuffle
    baseline due to Central Limit Theorem. This version uses global baseline
    for computational efficiency (~100x faster) with equivalent results.

IMPLEMENTATION STANDARDS:
    1. Sliding windows with 50% overlap (temporal resolution)
    2. Global baseline calculation (text's own distribution)
    3. True KL divergence formula: Σ p_k * log₂(p_k / π_k)
    4. Whole-word matching with word boundaries

USAGE:
    1. Set TEXT_FILE path (line 90)
    2. Modify motif_dict for your domain (lines 95-166)
    3. Run: python symbolic_entropy_streamlined.py
    
    Advanced: python symbolic_entropy_streamlined.py <path_to_text_file>

OUTPUTS:
    - <textname>_se_heatmap.png          (dual heatmap: raw density + KL)
    - <textname>_se_timeseries.png       (H, Σ, SE over windows)
    - <textname>_se_results.csv          (raw data for analysis)
    - Console: Statistical summary and diagnostics

DEPENDENCIES:
    - numpy
    - pandas
    - matplotlib
    - re, collections, random (standard library)
    
CITATION:
    Kurian, M. (2025). Symbolic Entropy: A Mathematical Framework for 
    Quantifying Meaning Density in Text. [Manuscript in preparation]

VERSION:
    1.1.0 - Streamlined (2025-01-14)
    Removed shuffle baseline (empirically equivalent to global for large texts)
    
LICENSE:
    [Specify license]

================================================================================
"""

import re
import os
import sys
import numpy as np
import pandas as pd
import random
import matplotlib.pyplot as plt
from collections import Counter
from scipy.signal import find_peaks

# Try to import docx for .docx file support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ============================================================================
# CONFIGURATION PARAMETERS
# ============================================================================

# Default text file (can be overridden via command line)
TEXT_FILE = "C:/Users/Michael Kurian/Desktop/Advanced Researches/TXT Files for SE/Fellowship of the RIng SE.txt"

# Random seed for reproducibility
RANDOM_SEED = 42

# ============================================================================
# MOTIF DICTIONARY - FELLOWSHIP OF THE RING
# ============================================================================
# Format: {Category: [list of lowercase single-token words]}
# These represent archetypal themes tracked through the narrative

motif_dict = {
    'The One Ring': [
        'ring', 'precious', 'gold', 'band', 'circle', 'master-ring', 'ruling', 
        'gollum', 'invisibility', 'chain', 'burden', 'magic', 'finger', 'pocket', 
        'vanish', 'secret', 'found', 'gave', 'took', 'kept', 'possess', 'will', 'power'
    ],
    
    'The Fellowship': [
        'fellowship', 'companions', 'company', 'quest', 'group', 'nine', 'walkers', 
        'unity', 'alliance', 'brotherhood', 'journey', 'friends', 'party', 'together', 
        'travel', 'set', 'band', 'walking', 'adventure',
        'frodo', 'sam', 'samwise', 'merry', 'meriadoc', 'pippin', 'peregrin',
        'aragorn', 'strider', 'legolas', 'gimli', 'boromir', 'gandalf'
    ],
    
    'The Shire': [
        'shire', 'hobbiton', 'bag', 'end', 'bywater', 'hobbit', 'green', 'hill', 
        'westfarthing', 'eastfarthing', 'southfarthing', 'northfarthing', 'home', 
        'garden', 'field', 'post', 'road', 'row', 'gaffer', 'hole', 'peaceful', 
        'comfort', 'folk'
    ],
    
    'The Road/Journey': [
        'road', 'journey', 'path', 'travel', 'wander', 'quest', 'adventure', 'trail', 
        'route', 'passage', 'walking', 'miles', 'crossing', 'errand', 'march', 'start', 
        'way', 'ahead', 'behind', 'go', 'leave', 'walk', 'step'
    ],
    
    'Light and Darkness': [
        'light', 'darkness', 'shadow', 'dark', 'bright', 'gloom', 'shining', 'night', 
        'dawn', 'dusk', 'sunlight', 'lantern', 'lamp', 'moon', 'star', 'glow', 'fire', 
        'sun', 'morning', 'evening', 'shine', 'shadowy', 'black', 'white', 'pale', 'stars'
    ],
    
    'The Shadow': [
        'shadow', 'sauron', 'enemy', 'black', 'rider', 'nazgul', 'ringwraith', 'wraith', 
        'darkness', 'evil', 'threat', 'eye', 'mordor', 'pursuit', 'fear', 'dread', 
        'cloak', 'hood', 'sniff', 'follow', 'hunt', 'search', 'danger', 'servant', 
        'master', 'power'
    ],
    
    'Nature and the Old Forest': [
        'forest', 'tree', 'old', 'woods', 'river', 'willow', 'grass', 'leaf', 'root', 
        'hedge', 'meadow', 'glade', 'bark', 'moss', 'stream', 'thicket', 'field', 
        'hill', 'water', 'wood', 'branch', 'earth', 'green', 'wild'
    ],
    
    'Songs and Poetry': [
        'song', 'singing', 'poem', 'verse', 'tune', 'chant', 'music', 'melody', 'rhyme', 
        'ballad', 'lay', 'recite', 'chorus', 'elven', 'hobbit', 'sing', 'voice', 'words', 
        'dance', 'laugh', 'cheer', 'merry'
    ],
    
    'Hospitality and Feasting': [
        'feast', 'supper', 'meal', 'party', 'table', 'food', 'drink', 'ale', 'breakfast', 
        'lunch', 'dinner', 'banquet', 'kitchen', 'cook', 'provision', 'toast', 'wine', 
        'eat', 'mug', 'jug', 'bread', 'beer', 'plate', 'dish', 'bottle', 'present', 
        'host', 'guest', 'welcome', 'invitation'
    ],
    
    'Gifts and Tokens': [
        'gift', 'token', 'present', 'cloak', 'sword', 'treasure', 'mathom', 'keepsake', 
        'heirloom', 'trinket', 'ring', 'blade', 'shield', 'necklace', 'relic', 'label', 
        'envelope', 'parcel', 'letter', 'gold', 'silver', 'spoon', 'pen', 'ink', 'key', 
        'chest', 'bag', 'package'
    ],
    
    'Hidden Identity and Disguise': [
        'disguise', 'hidden', 'invisibility', 'secret', 'cloak', 'underhill', 'mask', 
        'concealed', 'alias', 'shadowed', 'pseudonym', 'eavesdrop', 'spy', 'vanish', 
        'slip', 'unseen', 'invisible', 'pocket', 'name', 'stranger', 'unknown', 'hiding'
    ],
    
    'Friendship and Loyalty': [
        'friend', 'friendship', 'loyal', 'companion', 'trust', 'bond', 'fellowship', 
        'support', 'devotion', 'kinship', 'ally', 'camaraderie', 'faith', 'together', 
        'help', 'faithful', 'dear', 'close'
    ],
    
    'Ancient Lore and History': [
        'lore', 'history', 'legend', 'tale', 'story', 'ancient', 'elendil', 'myth', 
        'record', 'annals', 'memory', 'past', 'old', 'days', 'genealogy', 'saga', 
        'told', 'remember'
    ],
    
    'Guardians and Guides': [
        'guide', 'guardian', 'gandalf', 'elrond', 'tom', 'bombadil', 'protector', 
        'mentor', 'leader', 'watcher', 'helper', 'advisor', 'steward', 'teacher', 
        'wisdom', 'help', 'advice', 'counsel', 'lead', 'protect', 'watch', 'wise'
    ],
    
    'Temptation and Choice': [
        'temptation', 'choice', 'choose', 'decision', 'will', 'mercy', 'struggle', 
        'test', 'resolve', 'dilemma', 'trial', 'fate', 'free', 'crossroads', 'resist', 
        'desire', 'wish', 'decide', 'must', 'cannot', 'want', 'should', 'need', 'willpower'
    ]
}

# ============================================================================
# FILE LOADING FUNCTIONS
# ============================================================================

def load_text_file(filepath):
    """
    Load text from .txt or .docx file with smart encoding detection.
    
    Args:
        filepath: Path to text file (.txt) or Word document (.docx)
        
    Returns:
        str: Text content
    """
    _, ext = os.path.splitext(filepath)
    ext = ext.lower()
    
    if ext == '.docx':
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for .docx files. Install with: pip install python-docx")
        doc = Document(filepath)
        text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        print(f"✓ Loaded .docx file: {len(doc.paragraphs)} paragraphs")
        return text
    elif ext == '.txt':
        # Try common encodings
        for encoding in ['utf-8', 'utf-8-sig', 'cp1252', 'latin-1', 'iso-8859-1']:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    text = f.read()
                print(f"✓ Loaded with {encoding} encoding")
                return text
            except (UnicodeDecodeError, UnicodeError):
                continue
        # Last resort: utf-8 with error handling
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()
        print(f"⚠ Loaded with utf-8 (some characters replaced)")
        return text
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .txt or .docx")

# ============================================================================
# CORE ANALYSIS FUNCTIONS
# ============================================================================

def tokenize_text(text):
    """
    Tokenize text using word boundaries and lowercase normalization.
    
    Args:
        text (str): Raw text input
        
    Returns:
        list: List of lowercase tokens
    """
    text = re.sub(r'[^a-zA-Z\s\']', ' ', text.lower())
    tokens = re.findall(r'\b\w+(?:\'\w+)?\b', text)
    return tokens


def count_motifs_in_window(window_tokens, motif_dict):
    """
    Count motif occurrences in a window using whole-word matching.
    
    Args:
        window_tokens (list): List of tokens in current window
        motif_dict (dict): Dictionary of {category: [words]}
        
    Returns:
        dict: {category: count} for each motif category
    """
    counts = {category: 0 for category in motif_dict.keys()}
    window_set = set(window_tokens)
    
    for category, words in motif_dict.items():
        for word in words:
            if word in window_set:
                counts[category] += window_tokens.count(word)
    
    return counts


def calculate_global_baseline(tokens, motif_dict):
    """
    Calculate global baseline as proportion of each motif in full text.
    
    Formula: π_k = (total count of motif k) / (total tokens)
    
    For large texts (>50k tokens), this converges to shuffle baseline
    due to Central Limit Theorem, providing equivalent results with
    ~100x faster computation.
    
    Args:
        tokens (list): All tokens in text
        motif_dict (dict): Motif categories
        
    Returns:
        dict: {category: proportion} baseline values
    """
    N = len(tokens)
    baseline = {}
    
    for category, words in motif_dict.items():
        total_count = sum(tokens.count(word) for word in words)
        baseline[category] = total_count / N
    
    return baseline


def calculate_shannon_entropy(window_tokens):
    """
    Calculate Shannon entropy (H) for a window.
    
    Formula: H = -Σ p(x) * log₂(p(x))
    
    Args:
        window_tokens (list): Tokens in current window
        
    Returns:
        float: Shannon entropy in bits per token
    """
    if len(window_tokens) == 0:
        return 0.0
    
    word_counts = Counter(window_tokens)
    total = len(window_tokens)
    
    H = 0.0
    for count in word_counts.values():
        p = count / total
        if p > 0:
            H -= p * np.log2(p)
    
    return H


def calculate_sigma_kl(observed, baseline_proportions, window_size):
    """
    Calculate Sigma (Σ) using true KL divergence formula.
    
    Formula: Σ_KL = Σ p_k * log₂(p_k / π_k)
    
    Where:
        p_k = observed proportion of motif k in window
        π_k = baseline proportion of motif k
    
    Args:
        observed (dict): {category: count} in current window
        baseline_proportions (dict): {category: proportion} baseline
        window_size (int): Size of window for normalization
        
    Returns:
        float: KL divergence in bits per token
    """
    sigma_kl = 0.0
    
    for category in observed.keys():
        obs_count = observed[category]
        p_k = obs_count / window_size
        pi_k = baseline_proportions[category]
        
        if p_k > 0 and pi_k > 0:
            sigma_kl += p_k * np.log2(p_k / pi_k)
    
    return sigma_kl


# ============================================================================
# MAIN ANALYSIS PIPELINE
# ============================================================================

def run_se_analysis(text_path, motif_dict):
    """
    Execute complete Symbolic Entropy analysis pipeline.
    
    Args:
        text_path (str): Path to input text file
        motif_dict (dict): Motif category definitions
        
    Returns:
        tuple: (results_df, raw_densities, kl_contributions, motif_dict, 
                baseline, window_size, total_tokens)
    """
    print(f"\n{'='*70}")
    print(f"SYMBOLIC ENTROPY (SE) ANALYSIS")
    print(f"{'='*70}\n")
    
    # Validate file exists
    if not os.path.exists(text_path):
        raise FileNotFoundError(f"Text file not found: {text_path}")
    
    print(f"Loading text: {os.path.basename(text_path)}")
    text = load_text_file(text_path)
    
    # Tokenization
    print("Tokenizing...")
    tokens = tokenize_text(text)
    total_tokens = len(tokens)
    print(f"✓ Total tokens: {total_tokens:,}")
    
    # Display motif structure
    print(f"\n✓ Loaded {len(motif_dict)} motif categories:")
    for cat, words in motif_dict.items():
        print(f"  - {cat}: {len(words)} words")
    
    # Calculate window parameters (~110 windows with 50% overlap)
    target_windows = 110
    window_size = int(total_tokens / (1 + (target_windows - 1) / 2))
    step_size = window_size // 2  # MANDATORY: 50% overlap
    n_windows = (total_tokens - window_size) // step_size + 1
    
    print(f"\n{'='*70}")
    print(f"WINDOW PARAMETERS")
    print(f"{'='*70}")
    print(f"Window size: {window_size} tokens")
    print(f"Step size: {step_size} tokens (50% overlap)")
    print(f"Number of windows: {n_windows}")
    
    # Generate global baseline
    print(f"\nCalculating global baseline...")
    baseline = calculate_global_baseline(tokens, motif_dict)
    
    # Sample baseline values for diagnostic
    print(f"\nBaseline (first 3 categories):")
    for cat in list(motif_dict.keys())[:3]:
        print(f"  {cat}: {baseline[cat]:.6f}")
    
    # Window-by-window analysis
    print(f"\n{'='*70}")
    print(f"ANALYZING WINDOWS")
    print(f"{'='*70}\n")
    
    results = []
    raw_densities = {cat: [] for cat in motif_dict.keys()}
    kl_contributions = {cat: [] for cat in motif_dict.keys()}
    
    for i in range(0, total_tokens - window_size + 1, step_size):
        window = tokens[i:i + window_size]
        
        # Calculate H
        H = calculate_shannon_entropy(window)
        
        # Count motifs
        observed = count_motifs_in_window(window, motif_dict)
        
        # Calculate Σ
        sigma = calculate_sigma_kl(observed, baseline, window_size)
        
        # Calculate SE
        SE = H + sigma
        
        results.append({
            'window_index': len(results),
            'start_token': i,
            'H': H,
            'Sigma': sigma,
            'SE': SE
        })
        
        # Store visualization data
        for cat in motif_dict.keys():
            obs_count = observed[cat]
            
            # Raw density (thematic presence)
            p_k = obs_count / window_size
            raw_densities[cat].append(p_k)
            
            # KL contribution per category
            pi_k = baseline[cat]
            if p_k > 0 and pi_k > 0:
                kl_k = p_k * np.log2(p_k / pi_k)
            else:
                kl_k = 0.0
            kl_contributions[cat].append(kl_k)
        
        if (len(results)) % 20 == 0:
            print(f"  Processed window {len(results)}/{n_windows}")
    
    print(f"\n✓ Analysis complete: {len(results)} windows processed")
    
    results_df = pd.DataFrame(results)
    
    return (results_df, raw_densities, kl_contributions, motif_dict, 
            baseline, window_size, total_tokens, tokens)


# ============================================================================
# VISUALIZATION FUNCTIONS
# ============================================================================

def plot_dual_heatmap(results_df, raw_densities, kl_contributions, motif_dict, output_prefix):
    """
    Generate dual heatmap visualization (raw density + KL divergence).
    
    Args:
        results_df (DataFrame): Analysis results
        raw_densities (dict): Raw motif proportions
        kl_contributions (dict): KL divergence per category
        motif_dict (dict): Motif categories
        output_prefix (str): Prefix for output filename
    """
    print(f"\n{'='*70}")
    print(f"GENERATING HEATMAP")
    print(f"{'='*70}\n")
    
    categories = list(motif_dict.keys())
    raw_data = np.array([raw_densities[cat] for cat in categories])
    kl_data = np.array([kl_contributions[cat] for cat in categories])
    
    print(f"Raw data range: {raw_data.min():.4f} to {raw_data.max():.4f}")
    print(f"KL data range: {kl_data.min():.6f} to {kl_data.max():.6f}")
    
    # Create dual subplot figure
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(24, 10))
    
    # LEFT: Raw Density
    im1 = ax1.imshow(raw_data, aspect='auto', cmap='plasma', 
                     interpolation='nearest', origin='lower')
    
    # Extract and scale Sigma and H values for overlay
    sigma_values = results_df['Sigma'].values
    H_values = results_df['H'].values
    window_indices = results_df['window_index'].values
    
    # Scale Sigma to heatmap bounds
    if sigma_values.max() > sigma_values.min():
        sigma_scaled = ((sigma_values - sigma_values.min()) / 
                       (sigma_values.max() - sigma_values.min())) * (len(categories) - 1)
    else:
        sigma_scaled = np.zeros_like(sigma_values)
    
    # Scale H to heatmap bounds
    if H_values.max() > H_values.min():
        H_scaled = ((H_values - H_values.min()) / 
                   (H_values.max() - H_values.min())) * (len(categories) - 1)
    else:
        H_scaled = np.zeros_like(H_values)
    
    # Plot both lines on left axis
    ax1.plot(window_indices, sigma_scaled, color='white', linewidth=2.5, alpha=0.5)
    ax1.plot(window_indices, H_scaled, color='cyan', linewidth=2.5, alpha=0.5)
    ax1.set_ylim(-0.5, len(categories) - 0.5)
    ax1.set_xlabel('Window Index', fontsize=12)
    ax1.set_ylabel('Motif Category', fontsize=12)
    ax1.set_yticks(range(len(categories)))
    ax1.set_yticklabels(categories, fontsize=9)
    ax1.set_title('Method 1: RAW DENSITY\n(Thematic Presence - Where motifs appear)', 
                  fontsize=13, fontweight='bold')
    
    cbar1 = plt.colorbar(im1, ax=ax1)
    cbar1.set_label('Proportion of Window', fontsize=10)
    
    # RIGHT: KL Divergence
    im2 = ax2.imshow(kl_data, aspect='auto', cmap='plasma', 
                     interpolation='nearest', origin='lower')
    
    # Plot both lines on right axis
    ax2.plot(window_indices, sigma_scaled, color='white', linewidth=2.5, alpha=0.5)
    ax2.plot(window_indices, H_scaled, color='cyan', linewidth=2.5, alpha=0.5)
    ax2.set_ylim(-0.5, len(categories) - 0.5)
    ax2.set_xlabel('Window Index', fontsize=12)
    ax2.set_ylabel('Motif Category', fontsize=12)
    ax2.set_yticks(range(len(categories)))
    ax2.set_yticklabels(categories, fontsize=9)
    ax2.set_title('Method 2: KL DIVERGENCE (ΣKL)\n(Structural Surprise - Where motifs cluster)', 
                  fontsize=13, fontweight='bold')
    
    cbar2 = plt.colorbar(im2, ax=ax2)
    cbar2.set_label('KL Contribution (bits/token)', fontsize=10)
    
    fig.suptitle('Symbolic Entropy Analysis - Dual Method Visualization', 
                 fontsize=16, fontweight='bold', y=0.98)
    
    plt.tight_layout()
    filename = f'{output_prefix}_se_heatmap.png'
    plt.savefig(filename, dpi=300, bbox_inches='tight')
    print(f"✓ Saved: {filename}")
    plt.show()


def plot_timeseries(results_df, output_prefix):
    """
    Generate time series plots for H, Σ, and SE.
    
    Args:
        results_df (DataFrame): Analysis results
        output_prefix (str): Prefix for output filename
    """
    print(f"\n{'='*70}")
    print(f"GENERATING TIME SERIES")
    print(f"{'='*70}\n")
    
    fig, axes = plt.subplots(3, 1, figsize=(14, 10), sharex=True)
    
    # H plot
    axes[0].plot(results_df['window_index'], results_df['H'], 
                 color='steelblue', linewidth=1.5)
    axes[0].set_ylabel('H (bits/token)', fontsize=11)
    axes[0].set_title('Shannon Entropy (H) - Lexical Diversity', 
                      fontsize=12, fontweight='bold')
    axes[0].grid(True, alpha=0.3)
    
    # Σ plot
    axes[1].plot(results_df['window_index'], results_df['Sigma'], 
                 color='crimson', linewidth=1.5)
    axes[1].set_ylabel('Σ (bits/token)', fontsize=11)
    axes[1].set_title('Sigma (Σ) - Motif Concentration', 
                      fontsize=12, fontweight='bold')
    axes[1].grid(True, alpha=0.3)
    
    # SE plot
    axes[2].plot(results_df['window_index'], results_df['SE'], 
                 color='darkgreen', linewidth=1.5)
    axes[2].set_ylabel('SE (bits/token)', fontsize=11)
    axes[2].set_xlabel('Window Index', fontsize=11)
    axes[2].set_title('SE = H + Σ - Total Symbolic Entropy', 
                      fontsize=12, fontweight='bold')
    axes[2].grid(True, alpha=0.3)
    
    plt.tight_layout()
    filename = f'{output_prefix}_se_timeseries.png'
    plt.savefig(filename, dpi=300, bbox_inches='tight')
    print(f"✓ Saved: {filename}")
    plt.show()


# ============================================================================
# STATISTICAL SUMMARY
# ============================================================================

def print_publication_statistics(results_df, text_path, total_tokens, window_size, n_windows):
    """
    Print publication-ready statistical summary.
    """
    print(f"\n{'='*70}")
    print(f"PUBLICATION-READY STATISTICS")
    print(f"{'='*70}")
    print(f"Text: {os.path.basename(text_path)}")
    print(f"Total tokens: {total_tokens:,}")
    print(f"Window size: {window_size} tokens")
    print(f"Number of windows: {n_windows}")
    print(f"Motif categories: {len(motif_dict)}")
    print(f"")
    print(f"Shannon Entropy (H):")
    print(f"  Mean: {results_df['H'].mean():.4f} ± {results_df['H'].std():.4f} bits/token")
    print(f"  Range: [{results_df['H'].min():.4f}, {results_df['H'].max():.4f}]")
    print(f"")
    print(f"Sigma (Σ):")
    print(f"  Mean: {results_df['Sigma'].mean():.6f} ± {results_df['Sigma'].std():.6f} bits/token")
    print(f"  Range: [{results_df['Sigma'].min():.6f}, {results_df['Sigma'].max():.6f}]")
    print(f"")
    print(f"Symbolic Entropy (SE = H + Σ):")
    print(f"  Mean: {results_df['SE'].mean():.4f} ± {results_df['SE'].std():.4f} bits/token")
    print(f"  Range: [{results_df['SE'].min():.4f}, {results_df['SE'].max():.4f}]")
    print(f"")
    print(f"Note: Global baseline used (equivalent to shuffle for large texts)")
    print(f"      Falsifiability validated separately via text randomization")
    print(f"{'='*70}\n")




# ============================================================================
# PEAK AND VALLEY DETECTION WITH TEXT EXTRACTION
# ============================================================================

def detect_peaks_and_valleys(values, n_peaks=3, n_valleys=3, min_distance=5):
    """
    Detect top N peaks and deepest N valleys in a signal.
    
    Args:
        values (array): Signal values (typically Σ or SE)
        n_peaks (int): Number of peaks to identify
        n_valleys (int): Number of valleys to identify
        min_distance (int): Minimum windows between detected points
        
    Returns:
        tuple: (peaks_list, valleys_list) where each is [(idx, value), ...]
    """
    values_array = np.array(values)
    
    # Detect peaks (local maxima)
    peak_indices, _ = find_peaks(
        values_array,
        prominence=0.001,  # Must stand out from neighbors
        distance=min_distance
    )
    
    # Sort peaks by value, take top N
    peak_data = [(int(idx), float(values_array[idx])) for idx in peak_indices]
    peak_data.sort(key=lambda x: x[1], reverse=True)
    top_peaks = peak_data[:n_peaks]
    
    # Detect valleys (local minima) by inverting the signal
    valley_indices, _ = find_peaks(
        -values_array,
        prominence=0.001,
        distance=min_distance
    )
    
    # Sort valleys by value (ascending = deepest valleys first)
    valley_data = [(int(idx), float(values_array[idx])) for idx in valley_indices]
    valley_data.sort(key=lambda x: x[1])
    top_valleys = valley_data[:n_valleys]
    
    return top_peaks, top_valleys


def extract_window_text(window_idx, tokens, window_size, step_size, 
                        context_words=100, raw_text=None):
    """
    Extract text passage from a specific window.
    
    Args:
        window_idx (int): Window index
        tokens (list): All tokens in text
        window_size (int): Size of sliding window
        step_size (int): Step between windows
        context_words (int): Additional words to show before/after window
        raw_text (str): Original text for better formatting (optional)
        
    Returns:
        dict: {
            'window_idx': int,
            'window_text': str (core window),
            'full_context': str (with surrounding context),
            'start_position': int,
            'end_position': int
        }
    """
    # Calculate token indices for this window
    start_token = window_idx * step_size
    end_token = start_token + window_size
    
    # Get window tokens
    window_tokens = tokens[start_token:end_token]
    window_text = ' '.join(window_tokens)
    
    # Get extended context
    context_start = max(0, start_token - context_words)
    context_end = min(len(tokens), end_token + context_words)
    context_tokens = tokens[context_start:context_end]
    context_text = ' '.join(context_tokens)
    
    return {
        'window_idx': window_idx,
        'window_text': window_text,
        'full_context': context_text,
        'start_position': start_token,
        'end_position': end_token,
        'context_start': context_start,
        'context_end': context_end
    }


def analyze_window_motifs(window_tokens, motif_dict):
    """
    Identify which motifs appear in a window and their frequencies.
    
    Args:
        window_tokens (list): Tokens in the window
        motif_dict (dict): Motif dictionary
        
    Returns:
        dict: {category: count} sorted by count (descending)
    """
    motif_counts = {}
    
    for category, words in motif_dict.items():
        count = sum(1 for token in window_tokens if token in words)
        if count > 0:
            motif_counts[category] = count
    
    return dict(sorted(motif_counts.items(), key=lambda x: x[1], reverse=True))


def plot_peaks_and_valleys(results_df, tokens, window_size, step_size, 
                           motif_dict, output_prefix, raw_text=None):
    """
    Create visualization showing SE/Σ with annotated peaks and valleys,
    plus text excerpts from key moments.
    
    Args:
        results_df (DataFrame): SE analysis results
        tokens (list): All tokens
        window_size (int): Window size used
        step_size (int): Step size used
        motif_dict (dict): Motif dictionary
        output_prefix (str): Filename prefix
        raw_text (str): Original text (optional)
    """
    print(f"\n{'='*70}")
    print(f"GENERATING PEAK AND VALLEY ANALYSIS")
    print(f"{'='*70}\n")
    
    # Detect peaks and valleys in Σ
    sigma_values = results_df['Sigma'].values
    peaks, valleys = detect_peaks_and_valleys(sigma_values, n_peaks=3, n_valleys=3)
    
    print(f"Top 3 Peaks (Highest Σ):")
    for rank, (idx, val) in enumerate(peaks, 1):
        print(f"  Peak {rank}: Window {idx}, Σ = {val:.6f}")
    
    print(f"\nTop 3 Valleys (Lowest Σ):")
    for rank, (idx, val) in enumerate(valleys, 1):
        print(f"  Valley {rank}: Window {idx}, Σ = {val:.6f}")
    
    # Create figure with main plot + text excerpts
    fig = plt.figure(figsize=(16, 14))
    gs = fig.add_gridspec(4, 2, height_ratios=[3, 1, 1, 1], hspace=0.5, wspace=0.4)
    
    # Main plot: Σ over time with highlighted peaks and valleys
    ax_main = fig.add_subplot(gs[0, :])
    
    window_indices = results_df['window_index'].values
    
    # Plot Σ line
    ax_main.plot(window_indices, sigma_values, color='crimson', 
                 linewidth=2, label='Σ (Motif Concentration)', alpha=0.7)
    
    # Mark peaks with upward arrows
    for rank, (idx, val) in enumerate(peaks, 1):
        ax_main.scatter(idx, val, color='gold', s=300, marker='^', 
                       edgecolors='black', linewidths=2, zorder=5)
        ax_main.annotate(f'Peak {rank}', xy=(idx, val), 
                        xytext=(0, 20), textcoords='offset points',
                        ha='center', fontsize=11, fontweight='bold',
                        bbox=dict(boxstyle='round,pad=0.5', facecolor='gold', alpha=0.8),
                        arrowprops=dict(arrowstyle='->', lw=2))
    
    # Mark valleys with downward arrows
    for rank, (idx, val) in enumerate(valleys, 1):
        ax_main.scatter(idx, val, color='lightblue', s=300, marker='v', 
                       edgecolors='black', linewidths=2, zorder=5)
        ax_main.annotate(f'Valley {rank}', xy=(idx, val), 
                        xytext=(0, -20), textcoords='offset points',
                        ha='center', fontsize=11, fontweight='bold',
                        bbox=dict(boxstyle='round,pad=0.5', facecolor='lightblue', alpha=0.8),
                        arrowprops=dict(arrowstyle='->', lw=2))
    
    ax_main.set_xlabel('Window Index', fontsize=12, fontweight='bold')
    ax_main.set_ylabel('Σ (bits/token)', fontsize=12, fontweight='bold')
    ax_main.set_title('Symbolic Entropy: Key Moments (Top 3 Peaks & Valleys)', 
                     fontsize=14, fontweight='bold')
    ax_main.grid(True, alpha=0.3)
    ax_main.legend(loc='upper right', fontsize=10)
    
    # Text excerpt panels - simple 3x2 grid
    excerpt_axes = [
        (fig.add_subplot(gs[1, 0]), 'Peak 1', peaks[0] if len(peaks) > 0 else None, 'gold'),
        (fig.add_subplot(gs[1, 1]), 'Peak 2', peaks[1] if len(peaks) > 1 else None, 'gold'),
        (fig.add_subplot(gs[2, 0]), 'Peak 3', peaks[2] if len(peaks) > 2 else None, 'gold'),
        (fig.add_subplot(gs[2, 1]), 'Valley 1', valleys[0] if len(valleys) > 0 else None, 'lightblue'),
        (fig.add_subplot(gs[3, 0]), 'Valley 2', valleys[1] if len(valleys) > 1 else None, 'lightblue'),
        (fig.add_subplot(gs[3, 1]), 'Valley 3', valleys[2] if len(valleys) > 2 else None, 'lightblue'),
    ]
    
    for ax, label, point_data, bgcolor in excerpt_axes:
        ax.axis('off')
        
        if point_data is None:
            continue
            
        idx, val = point_data
        
        # Extract text
        text_data = extract_window_text(idx, tokens, window_size, step_size, 
                                       context_words=50, raw_text=raw_text)
        
        # Get window tokens for motif analysis
        start = text_data['start_position']
        end = text_data['end_position']
        window_tokens = tokens[start:end]
        
        # Analyze motifs
        top_motifs = analyze_window_motifs(window_tokens, motif_dict)
        top_3_motifs = list(top_motifs.items())[:3]
        motif_str = ', '.join([f"{cat}: {cnt}" for cat, cnt in top_3_motifs])
        
        # Format text excerpt - keep it SHORT to match header size
        excerpt = text_data['window_text']
        # Limit to ~200 chars for compact display
        if len(excerpt) > 200:
            excerpt = excerpt[:200] + '...'
        
        # Add line breaks for better wrapping (every ~50 chars)
        words = excerpt.split()
        lines = []
        current_line = []
        current_length = 0
        for word in words:
            if current_length + len(word) + 1 > 50:  # Force wrap at ~50 chars
                lines.append(' '.join(current_line))
                current_line = [word]
                current_length = len(word)
            else:
                current_line.append(word)
                current_length += len(word) + 1
        if current_line:
            lines.append(' '.join(current_line))
        wrapped_excerpt = '\n'.join(lines[:4])  # Max 4 lines
        
        # Display with compact formatting
        title_text = f"{label} (Win {idx}, Σ={val:.5f})"
        motif_text = f"Motifs: {motif_str}"
        
        # Title
        ax.text(0.5, 0.95, title_text, 
               transform=ax.transAxes, fontsize=10, fontweight='bold',
               ha='center', va='top',
               bbox=dict(boxstyle='round,pad=0.4', facecolor=bgcolor, alpha=0.5))
        
        # Motifs (smaller, under title)
        ax.text(0.5, 0.75, motif_text, 
               transform=ax.transAxes, fontsize=8,
               ha='center', va='top', style='italic',
               bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.7))
        
        # Text excerpt (compact, wrapped)
        ax.text(0.5, 0.50, wrapped_excerpt, 
               transform=ax.transAxes, fontsize=7,
               ha='center', va='top', family='monospace',
               bbox=dict(boxstyle='round,pad=0.4', facecolor='lightyellow', 
                        alpha=0.8, edgecolor='gray', linewidth=0.5))
    
    plt.tight_layout()
    filename = f'{output_prefix}_peaks_valleys.png'
    plt.savefig(filename, dpi=300, bbox_inches='tight')
    print(f"\n✓ Saved: {filename}")
    plt.show()
    
    # Export detailed text data to CSV
    export_peaks_valleys_csv(peaks, valleys, tokens, window_size, step_size, 
                             motif_dict, output_prefix)


def export_peaks_valleys_csv(peaks, valleys, tokens, window_size, step_size,
                             motif_dict, output_prefix):
    """
    Export peak and valley text excerpts to CSV for detailed analysis.
    """
    data = []
    
    # Add peaks
    for rank, (idx, val) in enumerate(peaks, 1):
        text_data = extract_window_text(idx, tokens, window_size, step_size, context_words=100)
        start = text_data['start_position']
        end = text_data['end_position']
        window_tokens = tokens[start:end]
        motifs = analyze_window_motifs(window_tokens, motif_dict)
        
        data.append({
            'type': 'Peak',
            'rank': rank,
            'window_idx': idx,
            'sigma_value': val,
            'excerpt': text_data['window_text'][:500],
            'top_motif_1': list(motifs.keys())[0] if len(motifs) > 0 else '',
            'top_motif_1_count': list(motifs.values())[0] if len(motifs) > 0 else 0,
            'top_motif_2': list(motifs.keys())[1] if len(motifs) > 1 else '',
            'top_motif_2_count': list(motifs.values())[1] if len(motifs) > 1 else 0,
            'top_motif_3': list(motifs.keys())[2] if len(motifs) > 2 else '',
            'top_motif_3_count': list(motifs.values())[2] if len(motifs) > 2 else 0,
        })
    
    # Add valleys
    for rank, (idx, val) in enumerate(valleys, 1):
        text_data = extract_window_text(idx, tokens, window_size, step_size, context_words=100)
        start = text_data['start_position']
        end = text_data['end_position']
        window_tokens = tokens[start:end]
        motifs = analyze_window_motifs(window_tokens, motif_dict)
        
        data.append({
            'type': 'Valley',
            'rank': rank,
            'window_idx': idx,
            'sigma_value': val,
            'excerpt': text_data['window_text'][:500],
            'top_motif_1': list(motifs.keys())[0] if len(motifs) > 0 else '',
            'top_motif_1_count': list(motifs.values())[0] if len(motifs) > 0 else 0,
            'top_motif_2': list(motifs.keys())[1] if len(motifs) > 1 else '',
            'top_motif_2_count': list(motifs.values())[1] if len(motifs) > 1 else 0,
            'top_motif_3': list(motifs.keys())[2] if len(motifs) > 2 else '',
            'top_motif_3_count': list(motifs.values())[2] if len(motifs) > 2 else 0,
        })
    
    df = pd.DataFrame(data)
    filename = f'{output_prefix}_peaks_valleys_text.csv'
    df.to_csv(filename, index=False)
    print(f"✓ Saved: {filename}")

# ============================================================================
# MAIN EXECUTION
# ============================================================================

if __name__ == "__main__":
    # Set random seed for reproducibility
    random.seed(RANDOM_SEED)
    np.random.seed(RANDOM_SEED)
    
    # Handle command-line argument
    if len(sys.argv) > 1:
        TEXT_FILE = sys.argv[1]
        print(f"Using command-line specified file: {TEXT_FILE}")
    
    # Check file exists
    if not os.path.exists(TEXT_FILE):
        print(f"ERROR: File not found: {TEXT_FILE}")
        print(f"Usage: python {sys.argv[0]} [path_to_text_file]")
        sys.exit(1)
    
    # Generate output prefix from filename
    output_prefix = os.path.splitext(os.path.basename(TEXT_FILE))[0]
    
    # Run main analysis
    try:
        (results_df, raw_densities, kl_contributions, motif_dict_used, 
         baseline, window_size, total_tokens, tokens) = run_se_analysis(TEXT_FILE, motif_dict)
    except Exception as e:
        print(f"ERROR during analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
    
    n_windows = len(results_df)
    
    # Print publication statistics
    print_publication_statistics(results_df, TEXT_FILE, total_tokens, window_size, n_windows)
    
    # Generate visualizations
    plot_dual_heatmap(results_df, raw_densities, kl_contributions, 
                     motif_dict_used, output_prefix)
    
    plot_timeseries(results_df, output_prefix)
    
    # Generate peak and valley visualization with text excerpts
    step_size = window_size // 2  # 50% overlap as per implementation
    plot_peaks_and_valleys(results_df, tokens, window_size, step_size,
                          motif_dict_used, output_prefix)
    
    # Save results
    csv_filename = f'{output_prefix}_se_results.csv'
    results_df.to_csv(csv_filename, index=False)
    print(f"\n✓ Results saved: {csv_filename}")
    
    # Final summary
    print(f"\n{'='*70}")
    print(f"ANALYSIS COMPLETE")
    print(f"{'='*70}")
    print(f"Outputs generated:")
    print(f"  - {output_prefix}_se_heatmap.png")
    print(f"  - {output_prefix}_se_timeseries.png")
    print(f"  - {output_prefix}_se_results.csv")
    print(f"  - {output_prefix}_peaks_valleys.png")
    print(f"  - {output_prefix}_peaks_valleys_text.csv")
    print(f"\nFor publication, cite:")
    print(f"  Kurian, M. (2025). Symbolic Entropy: A Mathematical Framework")
    print(f"  for Quantifying Meaning Density in Text.")
    print(f"\nMethodological note:")
    print(f"  This version uses global baseline for computational efficiency.")
    print(f"  Empirical testing shows equivalence to shuffle baseline for")
    print(f"  large texts (>50k tokens) due to Central Limit Theorem.")
    print(f"  Falsifiability validated via separate text randomization tests.")
    print(f"{'='*70}\n")